#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_to_docx.py — minimal Markdown → docx exporter for paper.md.

Handles:
  - H1/H2/H3 headings
  - paragraphs (normal text, with bold **...** and italic *...*)
  - tables (pipe tables)
  - blockquotes / paragraphs starting with "> "
  - inline math $...$ and display math $$...$$
  - bullet lists "- ..."
  - code fences ``` ```
"""
import os, sys, re, argparse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_runs(p, text):
    """Add text into paragraph p, handling **bold** and *italic*."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|\$[^$\n]+\$)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); r.italic = True
        elif part.startswith("$") and part.endswith("$"):
            # inline math — rendered as italic Cambria-Math style
            r = p.add_run(part[1:-1])
            r.italic = True
            r.font.name = "Cambria Math"
        else:
            p.add_run(part)

def md_to_docx(md_path, docx_path):
    text = open(md_path, encoding="utf-8").read()
    # Preprocess display equations $$...$$
    lines = text.split("\n")

    doc = Document()
    # set font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    i = 0
    in_table = False
    table_rows = []
    in_code = False

    def flush_table():
        nonlocal table_rows
        if not table_rows: return
        # filter alignment row "|---|---|"
        table_rows = [r for r in table_rows if not re.match(r"^\s*\|?[\s:|-]+\|?\s*$", r)]
        if not table_rows: return
        cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in table_rows]
        ncols = max(len(r) for r in cells)
        cells = [r + [""]*(ncols-len(r)) for r in cells]
        t = doc.add_table(rows=len(cells), cols=ncols)
        t.style = "Light Grid Accent 1"
        for ri, row in enumerate(cells):
            for ci, val in enumerate(row):
                cell = t.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs(p, val)
                if ri == 0:
                    for run in p.runs:
                        run.bold = True
        doc.add_paragraph()
        table_rows = []

    while i < len(lines):
        ln = lines[i]
        stripped = ln.strip()

        # code fence
        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph(ln)
            p.style = doc.styles["Normal"]
            for r in p.runs: r.font.name = "Courier New"; r.font.size = Pt(9)
            i += 1
            continue

        # table line
        if stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:-1]:
            table_rows.append(ln); i += 1; continue
        else:
            if table_rows:
                flush_table()

        # display math $$...$$ (single line or block)
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[2:-2])
            r.italic = True; r.font.name = "Cambria Math"; r.font.size = Pt(12)
            i += 1
            continue
        if stripped == "$$":
            # block until next $$
            buf = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                buf.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("\n".join(buf))
            r.italic = True; r.font.name = "Cambria Math"; r.font.size = Pt(12)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)", ln)
        if m:
            level = len(m.group(1))
            head = m.group(2)
            p = doc.add_paragraph()
            p.style = doc.styles[f"Heading {min(level,4)}"]
            add_runs(p, head)
            i += 1
            continue

        # bullet
        if re.match(r"^\s*[-*]\s+", ln):
            content = re.sub(r"^\s*[-*]\s+", "", ln)
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, content)
            i += 1
            continue

        # numbered list
        if re.match(r"^\s*\d+\.\s+", ln):
            content = re.sub(r"^\s*\d+\.\s+", "", ln)
            p = doc.add_paragraph(style="List Number")
            add_runs(p, content)
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            content = re.sub(r"^>\s*", "", stripped)
            r = p.add_run(content); r.italic = True
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # blank line
        if not stripped:
            i += 1
            continue

        # plain paragraph (concat continuation lines)
        buf = [ln.rstrip()]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip(): break
            if nxt.lstrip().startswith(("#","|","- ","* ","> ","```","$$")): break
            if re.match(r"^\s*\d+\.\s+", nxt): break
            buf.append(nxt.rstrip())
            i += 1
        para = " ".join(buf).strip()
        p = doc.add_paragraph()
        add_runs(p, para)

    if table_rows:
        flush_table()

    doc.save(docx_path)

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("md")
    ap.add_argument("docx")
    args = ap.parse_args()
    md_to_docx(args.md, args.docx)
    print(f"OK -> {args.docx}  ({os.path.getsize(args.docx)/1024:.1f} KB)")
